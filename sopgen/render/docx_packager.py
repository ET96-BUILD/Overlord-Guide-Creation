"""Render a validated SOP document as a Word (.docx) file — text only.

Intended for non-technical users without Cowork: open the docx, eyeball
the procedure, copy-paste text into the Overlord / iFixit guide form.
Images are NOT embedded — users insert their own screenshots from the
sibling ``images/`` directory in the iFixit form, and the docx exists
purely as a readable text outline of the SOP.

Bullet styling cues map to iFixit's bullet styles so a reviewer can
apply the matching style on paste-in:

    Caution: ...    -> red text
    Note: ...       -> italic
    Reminder: ...   -> bold
    (default)       -> normal
"""

from __future__ import annotations

import logging
from pathlib import Path

from docx import Document
from docx.shared import RGBColor

logger = logging.getLogger(__name__)

# ── Style constants ─────────────────────────────────────────────────────
_CAUTION_RED = RGBColor(0xC0, 0x00, 0x00)

_PREFIX_CAUTION = "Caution:"
_PREFIX_NOTE = "Note:"
_PREFIX_REMINDER = "Reminder:"


# ── Public API ──────────────────────────────────────────────────────────


def write_docx(sop_dict: dict, output_path: Path) -> Path:
    """Write *sop_dict* as a text-only Word document at *output_path*.

    Parameters
    ----------
    sop_dict : dict
        The validated SOP document (matches ``SOPDocument.model_dump()``).
    output_path : Path
        Where to write the ``.docx`` file. Created (overwritten).
    """
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    # ── Title + intro ───────────────────────────────────────────────
    doc.add_heading(sop_dict.get("title", ""), level=1)
    if sop_dict.get("intro"):
        doc.add_paragraph(sop_dict["intro"])

    # ── Steps ───────────────────────────────────────────────────────
    steps = sorted(sop_dict.get("steps", []), key=lambda s: s.get("step_number", 0))
    for step in steps:
        _render_step(doc, step)

    # ── Warnings (only if non-empty) ────────────────────────────────
    warnings = sop_dict.get("warnings") or []
    if warnings:
        doc.add_heading("Warnings", level=2)
        for w in warnings:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(w)
            run.italic = True

    doc.save(str(output_path))
    logger.info("DOCX saved → %s", output_path)
    return output_path


# ── Per-step rendering ──────────────────────────────────────────────────


def _render_step(doc: Document, step: dict) -> None:
    n = step.get("step_number")
    title = step.get("step_title", "")
    doc.add_heading(f"Step {n}: {title}", level=2)

    for substep in step.get("substeps", []):
        _render_substep(doc, substep)


def _render_substep(doc: Document, substep: str) -> None:
    """Add one bulleted substep, styling it by its prefix word."""
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(substep)

    stripped = substep.lstrip()
    if stripped.startswith(_PREFIX_CAUTION):
        run.font.color.rgb = _CAUTION_RED
    elif stripped.startswith(_PREFIX_NOTE):
        run.italic = True
    elif stripped.startswith(_PREFIX_REMINDER):
        run.bold = True
    # else: default formatting
