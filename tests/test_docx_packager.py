"""Tests for the text-only .docx side-output."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from sopgen.render.docx_packager import write_docx


# ── Fixtures ────────────────────────────────────────────────────────────


def _make_sop(
    *,
    warnings: list[str] | None = None,
    extra_substeps: list[str] | None = None,
) -> dict:
    """Minimal valid SOP dict, mirroring SOPDocument.model_dump() shape."""
    base_substeps = ["Open the dashboard", "Click 'New Bill'"]
    if extra_substeps:
        base_substeps = base_substeps + extra_substeps
    return {
        "title": "How to Pay EAL Commission",
        "intro": "This procedure pays EAL their monthly commission before the 10th.",
        "settings": {"max_substeps_per_step": 4, "min_images_per_step": 1},
        "steps": [
            {
                "step_number": 1,
                "step_title": "Open the bill form",
                "substeps": base_substeps,
                "evidence": {
                    "recommended_screenshot_timestamps": ["00:05"],
                    "supporting_timestamps": [],
                },
                "images": [
                    {
                        "image_id": "step_1_img_1",
                        "filename": "step_1_img_1.png",
                        "url": "images/step_1_img_1.png",
                        "caption": "Bill form populated with vendor and posting period",
                    }
                ],
            },
            {
                "step_number": 2,
                "step_title": "Save and confirm",
                "substeps": ["Click Save", "Verify the new bill number"],
                "evidence": {
                    "recommended_screenshot_timestamps": ["00:30"],
                    "supporting_timestamps": [],
                },
                "images": [
                    {
                        "image_id": "step_2_img_1",
                        "filename": "step_2_img_1.png",
                        "url": "images/step_2_img_1.png",
                        "caption": "Save confirmation showing bill number 12345",
                    }
                ],
            },
        ],
        "warnings": warnings if warnings is not None else [],
    }


# ── Helpers for inspecting the generated doc ────────────────────────────


def _headings(doc: Document, level: int) -> list[str]:
    style_name = f"Heading {level}"
    return [p.text for p in doc.paragraphs if p.style.name == style_name]


def _all_text(doc: Document) -> str:
    return "\n".join(p.text for p in doc.paragraphs)


# ── File creation ───────────────────────────────────────────────────────


class TestFileCreation:
    def test_write_docx_produces_file_at_expected_path(self, tmp_path):
        out = tmp_path / "sop.docx"
        result = write_docx(_make_sop(), out)
        assert result == out
        assert out.exists()
        assert out.stat().st_size > 0

    def test_creates_parent_directory_if_missing(self, tmp_path):
        out = tmp_path / "nested" / "dir" / "sop.docx"
        write_docx(_make_sop(), out)
        assert out.exists()


# ── Document structure ──────────────────────────────────────────────────


class TestDocumentStructure:
    def test_title_is_heading_1(self, tmp_path):
        sop = _make_sop()
        out = tmp_path / "sop.docx"
        write_docx(sop, out)

        doc = Document(str(out))
        assert sop["title"] in _headings(doc, 1)

    def test_each_step_title_is_heading_2(self, tmp_path):
        sop = _make_sop()
        out = tmp_path / "sop.docx"
        write_docx(sop, out)

        doc = Document(str(out))
        h2s = _headings(doc, 2)
        for step in sop["steps"]:
            expected = f"Step {step['step_number']}: {step['step_title']}"
            assert expected in h2s, f"missing H2 for {expected!r}; got {h2s}"

    def test_intro_appears_in_body(self, tmp_path):
        sop = _make_sop()
        out = tmp_path / "sop.docx"
        write_docx(sop, out)
        doc = Document(str(out))
        assert sop["intro"] in _all_text(doc)

    def test_steps_rendered_in_step_number_order(self, tmp_path):
        sop = _make_sop()
        # Shuffle to confirm the renderer sorts by step_number.
        sop["steps"] = list(reversed(sop["steps"]))
        out = tmp_path / "sop.docx"
        write_docx(sop, out)

        doc = Document(str(out))
        h2s = _headings(doc, 2)
        idx1 = next(i for i, h in enumerate(h2s) if h.startswith("Step 1:"))
        idx2 = next(i for i, h in enumerate(h2s) if h.startswith("Step 2:"))
        assert idx1 < idx2


# ── Substeps + prefix preservation ──────────────────────────────────────


class TestSubsteps:
    def test_every_substep_appears_in_body(self, tmp_path):
        sop = _make_sop()
        out = tmp_path / "sop.docx"
        write_docx(sop, out)
        doc = Document(str(out))
        body = _all_text(doc)
        for step in sop["steps"]:
            for substep in step["substeps"]:
                assert substep in body, f"substep missing: {substep!r}"

    def test_caution_note_reminder_prefixes_preserved_verbatim(self, tmp_path):
        sop = _make_sop(
            extra_substeps=[
                "Caution: this submission cannot be undone",
                "Note: the posting period must be the prior month",
                "Reminder: complete this before the 10th",
            ]
        )
        out = tmp_path / "sop.docx"
        write_docx(sop, out)
        doc = Document(str(out))
        body = _all_text(doc)

        assert "Caution: this submission cannot be undone" in body
        assert "Note: the posting period must be the prior month" in body
        assert "Reminder: complete this before the 10th" in body

    def test_caution_substep_is_styled_red(self, tmp_path):
        sop = _make_sop(extra_substeps=["Caution: irreversible"])
        out = tmp_path / "sop.docx"
        write_docx(sop, out)
        doc = Document(str(out))

        caution_runs = [
            run
            for p in doc.paragraphs
            for run in p.runs
            if run.text.startswith("Caution:")
        ]
        assert caution_runs, "no Caution: run found"
        assert any(
            r.font.color.rgb is not None and str(r.font.color.rgb) == "C00000"
            for r in caution_runs
        )


# ── Text-only invariants (no embedded images / no captions) ─────────────


class TestTextOnly:
    def test_image_captions_are_not_embedded_in_body(self, tmp_path):
        """Captions describe screenshots — they belong with the screenshot,
        not in a text-only outline. Verify they're omitted entirely."""
        sop = _make_sop()
        out = tmp_path / "sop.docx"
        write_docx(sop, out)
        doc = Document(str(out))
        body = _all_text(doc)

        # Captions from the fixture must not leak into the docx
        for step in sop["steps"]:
            for img in step["images"]:
                assert img["caption"] not in body, (
                    f"caption leaked into text-only docx: {img['caption']!r}"
                )

    def test_no_image_xml_present(self, tmp_path):
        """Sanity: a text-only doc has no <pic:pic> picture frames."""
        sop = _make_sop()
        out = tmp_path / "sop.docx"
        write_docx(sop, out)
        doc = Document(str(out))

        # Iterate inline shapes (embedded pictures) — must be empty.
        assert len(doc.inline_shapes) == 0

    def test_signature_does_not_take_image_dir(self):
        """write_docx is text-only and accepts only (sop_dict, output_path)."""
        import inspect

        sig = inspect.signature(write_docx)
        params = list(sig.parameters.keys())
        assert "image_dir" not in params
        assert params == ["sop_dict", "output_path"]


# ── Warnings section ────────────────────────────────────────────────────


class TestWarningsSection:
    def test_no_warnings_heading_when_warnings_empty(self, tmp_path):
        out = tmp_path / "sop.docx"
        write_docx(_make_sop(warnings=[]), out)
        doc = Document(str(out))
        assert "Warnings" not in _headings(doc, 2)

    def test_warnings_heading_and_bullets_when_nonempty(self, tmp_path):
        warnings = [
            "Recording cuts off at 02:15 before the final confirmation",
            "Account number was redacted in the recording",
        ]
        out = tmp_path / "sop.docx"
        write_docx(_make_sop(warnings=warnings), out)
        doc = Document(str(out))

        assert "Warnings" in _headings(doc, 2)
        body = _all_text(doc)
        for w in warnings:
            assert w in body
